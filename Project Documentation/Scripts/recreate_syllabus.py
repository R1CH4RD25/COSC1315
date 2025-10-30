"""
Script to recreate the COSC 1315 Syllabus as a DOCX file
Using python-docx library for professional document creation
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_with_style(doc, text, level=1):
    """Add a heading with custom styling"""
    heading = doc.add_heading(text, level=level)
    heading.style.font.color.rgb = RGBColor(0, 0, 0)
    return heading

def add_table_cell_shading(cell, color):
    """Add shading/background color to a table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_course_info_table(doc):
    """Create the course information table at the top"""
    table = doc.add_table(rows=13, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Course information data
    info = [
        ('Course:', 'COSC 1315.35 - Introduction to Computer Programming'),
        ('Term:', 'FALL 2025'),
        ('Teacher:', 'Richard Sullivan'),
        ('Location:', 'Dual Credit - Woodson H.S.'),
        ('Credit Hours:', '3 Credit Hours'),
        ('Prerequisite:', 'None'),
        ('Class Hours:', 'Mon-Thurs - 10:45 AM to 11:42 AM'),
        ('Class Location:', 'Room 1204'),
        ('Class Style:', 'In Person'),
        ('Office Hours:', '9:42 AM - 10:42 AM (Mon - Thurs)'),
        ('Office Phone:', '940.204.6521 x1204'),
        ('Email:', 'richard.sullivan@cisco.edu'),
        ('Response time:', 'Within 24 Business Hours'),
    ]
    
    for i, (label, value) in enumerate(info):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        # Make labels bold
        row.cells[0].paragraphs[0].runs[0].font.bold = True
    
    return table

def create_grading_table(doc):
    """Create the grading policy table"""
    # Left table - Assignment weights
    table1 = doc.add_table(rows=4, cols=2)
    table1.style = 'Light Grid Accent 1'
    
    weights = [
        ('Programming Assignments', '65%'),
        ('Homework assignments', '25%'),
        ('Exam', '10%'),
        ('Total', '100%'),
    ]
    
    for i, (item, weight) in enumerate(weights):
        row = table1.rows[i]
        row.cells[0].text = item
        row.cells[1].text = weight
        if i == 3:  # Total row
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[1].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()  # Spacing
    
    # Right table - Grade scale
    p = doc.add_paragraph('Percentage Points Grade', style='Heading 3')
    table2 = doc.add_table(rows=5, cols=2)
    table2.style = 'Light Grid Accent 1'
    
    grades = [
        ('90 percent and above', 'A'),
        ('80 to 89.9', 'B'),
        ('70 to 79.9', 'C'),
        ('60 to 69.9', 'D'),
        ('Under 60', 'F'),
    ]
    
    for i, (range_text, grade) in enumerate(grades):
        row = table2.rows[i]
        row.cells[0].text = f"▪ {range_text}"
        row.cells[1].text = grade

def create_schedule_table(doc):
    """Create the course schedule table"""
    table = doc.add_table(rows=17, cols=4)  # 1 header + 16 data rows
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header = table.rows[0]
    header.cells[0].text = 'Week'
    header.cells[1].text = 'Dates'
    header.cells[2].text = 'Assignment / Project'
    header.cells[3].text = 'Due Date'
    
    # Make header bold
    for cell in header.cells:
        cell.paragraphs[0].runs[0].font.bold = True
        add_table_cell_shading(cell, 'D3D3D3')
    
    # Schedule data
    schedule = [
        ('0', '8/25–8/28', 'Introduction Week – Syllabus / Setup / Orientation', 'Thu, 8/28/25 @ 4:12 PM'),
        ('1', '9/2–9/4', 'Lesson 01 – Your First Python Program (6:08–8:08)\nLesson 01 – Quiz', 'Thu, 9/4/25 @ 4:12 PM'),
        ('2', '9/8–9/11', 'Lesson 02 – How Python Code Gets Executed (8:08–11:20)\nLesson 02 – Quiz', 'Thu, 9/11/25 @ 4:12 PM'),
        ('3', '9/15–9/18', 'Lesson 03 – Variables (12:56–18:16)\nLesson 04 – Receiving Input (18:16–22:08)\nLessons 03–04 – Quiz', 'Thu, 9/18/25 @ 4:12 PM'),
        ('4', '9/22–9/25', 'Lesson 05 – Type Conversion (22:40–29:28)\nLesson 05 – Quiz', 'Thu, 9/25/25 @ 4:12 PM'),
        ('5', '9/29–10/2', 'Lesson 06 – Strings (29:28–37:28)\nLesson 07 – Formatted Strings (37:28–40:48) & String Methods (40:48–48:32)\nLessons 06–07 – Quiz', 'Thu, 10/2/25 @ 4:12 PM'),
        ('6', '10/6–10/9', 'Lesson 08 – Arithmetic Operations (48:32–51:28), Operator Precedence (51:28–54:56), Math Functions (54:56–58:16)\nLesson 08 – Quiz', 'Thu, 10/9/25 @ 4:12 PM'),
        ('7', '10/13–10/16', 'Lesson 09 – If Statements (58:16–1:06:24)\nLesson 10 – Logical Operators (1:06:24–1:11:20) & Comparison Operators (1:11:20–1:16:16)\nLessons 09–10 – Quiz', 'Thu, 10/16/25 @ 4:12 PM'),
        ('8', '10/20–10/23', 'Lesson 11 – Weight Converter Program (1:16:16–1:20:40)\nLesson 11 – Quiz', 'Thu, 10/23/25 @ 4:12 PM'),
        ('9', '10/27–10/30', 'Lesson 12 – While Loops (1:20:40–1:41:44)\nLesson 12 – Quiz', 'Thu, 10/30/25 @ 4:12 PM'),
        ('10', '11/3–11/6', 'Lesson 13 – For Loops (1:41:44–1:47:44)\nLesson 14 – Nested Loops (1:47:44–1:55:44)\nLessons 13–14 – Quiz', 'Thu, 11/6/25 @ 4:12 PM'),
        ('11', '11/10–11/13', 'Lesson 15 – Lists (1:55:44–2:01:44)\nLesson 16 – 2D Lists (2:01:44–2:05:04)\nLessons 15–16 – Quiz', 'Thu, 11/13/25 @ 4:12 PM'),
        ('12', '11/17–11/20', 'Lesson 17 – List Methods (2:05:52–2:13:20)\nLesson 17 – Quiz', 'Thu, 11/20/25 @ 4:12 PM'),
        ('—', '11/24–11/27', 'No Class – Thanksgiving Break', '—'),
        ('13', '12/1–12/4', 'Lesson 18 – Tuples (2:13:20–2:15:28)\nLesson 18 – Quiz', 'Thu, 12/4/25 @ 4:12 PM'),
        ('14', '12/8–12/11', 'Finals Week – Comprehensive Review & Exam (no new lessons)', '—'),
    ]
    
    for i, (week, dates, assignment, due) in enumerate(schedule, 1):
        row = table.rows[i]
        row.cells[0].text = week
        row.cells[1].text = dates
        row.cells[2].text = assignment
        row.cells[3].text = due

def main():
    """Main function to create the syllabus"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ========== PAGE 1 ==========
    # Course Information Table
    create_course_info_table(doc)
    
    doc.add_paragraph()  # Spacing
    
    # Terms of Use
    add_heading_with_style(doc, 'Terms of Use', level=2)
    doc.add_paragraph(
        "A student's enrollment in Cisco College & Woodson ISD and within this course signifies "
        "acknowledgment of and agreement with the statements, disclaimers, policies, and procedures "
        "outlined within this syllabus and elsewhere in the Woodson ISD Parent/Student Handbook and "
        "the Cisco College Student Handbook. This Syllabus is a dynamic document. Elements of the "
        "course structure (e.g., dates and topics covered, but not policies) may be changed at the "
        "discretion of the instructor."
    )
    
    # Mission Statement
    add_heading_with_style(doc, 'Mission Statement', level=2)
    doc.add_paragraph(
        "The mission of Cisco College and Woodson ISD is to provide high-quality education through "
        "the use and implementation of proven teaching strategies and by utilizing 21st Century "
        "technology standards to implement an effective learner-driven environment to assist learners "
        "in areas of growth and skill development. We accomplish this through an emphasis on excellence "
        "in teaching, which is strengthened by experienced faculty and supported by outstanding administration."
    )
    
    doc.add_page_break()
    
    # ========== PAGE 2 ==========
    # Learning Objectives
    add_heading_with_style(doc, 'Learning Objectives', level=2)
    p = doc.add_paragraph(
        "Educators seek to prepare students with professional growth and advancement via key "
        "future-focused instruction. Students are provided with high-quality education to prepare them for "
        "various avenues of success after school. Students are offered degree programs to gain advanced dual "
        "credit hours in correlation with Cisco College. In addition, these institutions offer certification "
        "pathways for students aimed at fostering their learning goals and objectives."
    )
    
    doc.add_paragraph("The key learning objectives are as follows:", style='Heading 3')
    objectives = ['Leadership', 'Communication', 'Critical Thinking', 'Core Subject Mastery', 
                  'Hands-On Learning', 'Proficiency in Technology']
    for obj in objectives:
        doc.add_paragraph(f'● {obj}')
    
    # Course Description
    add_heading_with_style(doc, 'Course Description', level=2)
    doc.add_paragraph(
        "Introduction to computer programming for solving a variety of problems. This course is intended "
        "for non-computer science and non-computer engineering majors. Emphasis on the fundamentals of "
        "design, development, testing, implementation, and documentation of computer programs. Includes "
        "problem-solving with structured techniques and algorithms using pseudocode and/or graphical "
        "representations. The aim is to provide students with an understanding of the role computation can "
        "play in solving problems and to help students, regardless of their major, feel justifiably confident "
        "of their ability to write small programs that allow them to accomplish useful goals."
    )
    doc.add_paragraph()
    doc.add_paragraph("Programming assignments will use the latest Python programming language.").italic = True
    
    # Course Objectives
    add_heading_with_style(doc, 'Course Objectives', level=2)
    course_objs = [
        'Design and develop algorithms to solve problems.',
        'Demonstrate a fundamental understanding of software development methodologies, such as modular design, pseudo code, flowcharting, and structure charts.',
        'Demonstrate appropriate design, coding, testing, debugging, and documenting of computer programs that implement problem specifications and requirements.',
        'Apply computer programming concepts to new problems or situations.'
    ]
    for obj in course_objs:
        doc.add_paragraph(f'● {obj}')
    
    doc.add_page_break()
    
    # ========== PAGE 3 ==========
    # Core Objectives
    add_heading_with_style(doc, 'Core Objective (THECB)', level=2)
    doc.add_paragraph(
        "This course meets the following Core Objectives required by the Texas Higher Education Coordinating Board. "
        "Assessment may be based on the following: exams, post-tests, finals, assignments, and labs."
    )
    
    core_objs = [
        'Critical Thinking Skills (CT) - creative thinking, innovation, inquiry, analysis, evaluation, and synthesis of information',
        'Communication Skills (COM) - effective development, interpretation, and expression of ideas through written, oral, and visual communication',
        'Empirical and Quantitative Skills (EQS) - manipulation and analysis of numerical data or observable facts resulting in informed conclusions'
    ]
    for obj in core_objs:
        doc.add_paragraph(f'● {obj}')
    
    # Course Materials
    add_heading_with_style(doc, 'Course Materials (Including text, calculator, internet connectivity, software, virtual programs, etc.)', level=2)
    
    # Create materials table
    materials_table = doc.add_table(rows=1, cols=2)
    materials_table.style = 'Light Grid Accent 1'
    
    header = materials_table.rows[0]
    header.cells[0].text = 'Required Textbook'
    header.cells[1].text = 'Online Resources'
    for cell in header.cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    row = materials_table.add_row()
    row.cells[0].text = 'Introduction to Programming in Python 1st Edition\nRobert Sedgewick, Kevin Wayne, Robert Dondero\nAddison-Wesley Professional'
    row.cells[1].text = 'Google Colab\nCode with Mosh\nCodeHS\nFreecodecamp'
    
    doc.add_paragraph()
    p = doc.add_paragraph('Accessibility and Privacy Policies:')
    p.add_run('\nhttps://ciscocollege.instructure.com/courses/14725/pages/accessibility-and-privacy-for-course-technologies').font.color.rgb = RGBColor(0, 0, 255)
    
    # Course Grading Policies
    add_heading_with_style(doc, 'Course Grading Policies', level=2)
    create_grading_table(doc)
    
    doc.add_paragraph()
    doc.add_paragraph(
        "Final grades will be rounded to the higher grade if within 0.5% (i.e. 89.5, 79.5, etc. would be rounded up the higher grade)"
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Assignments are due no later than 2 weeks from the date the assignment is issued."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Dual credit classes will be cancelled in the event of Cisco College or ISD campus closure. "
        "Make-up work may be required. ISD campus events cannot disrupt dual credit classes."
    )
    
    doc.add_page_break()
    
    # ========== PAGE 4-5 ==========
    # Course Assignment Policies
    add_heading_with_style(doc, 'Course Assignment, Examination, and/or Project Policies', level=2)
    
    policies = [
        'Programming projects are a significant component of this course. These projects will be completed outside of class and students should be prepared to allocate several hours each week to this work.',
        'Students missing the exam due to leaving campus prior to a holiday or the end of the semester will receive an exam grade of F.',
        'The use of cell phones and other communication devices during an exam is not permitted and will result in an exam grade of F.',
        'Exams will be considered finished if a student leaves the room during the exam. A student may not resume work on an exam after leaving the exam room.',
        'FINAL EXAM DATE - Tuesday, December 9th'
    ]
    for policy in policies:
        doc.add_paragraph(f'● {policy}')
    
    doc.add_paragraph()
    doc.add_paragraph("Further details for each assignment will be provided in class or online via LMS.")
    doc.add_paragraph()
    
    # Course Schedule
    create_schedule_table(doc)
    
    doc.add_page_break()
    
    # ========== PAGE 6-7 ==========
    # Attendance
    add_heading_with_style(doc, 'Attendance', level=2)
    doc.add_paragraph(
        "Prompt and regular class attendance is considered necessary for satisfactory work. Attendance is "
        "defined by physical attendance or participation in an academically related activity. See the College "
        "Catalog for the full class attendance policy."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Cisco College recognizes that absence from class may occur due to illness, death or illness in the "
        "immediate family, observance of a religious holiday, or participation in a college-sponsored activity. "
        "(Absences due to participation in a college-sponsored activity must be authorized by the Vice President "
        "of Instruction.) When absences occur due to the above-stated reasons, the student is allowed to make up "
        "work missed; the professor may require the work to be made up within two weeks from its original due date. "
        "During a regular Fall or Spring semester, the following requirements apply for face-to-face classes. For a "
        "class that meets three times per week, a student is allowed six absences."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "For a class that meets two times per week, a student is allowed four absences. For a class that meets one "
        "time per week, a student is allowed two absences. If a student misses one more than the allowed number of "
        "absences, he/she may be dropped from the class if the professor deems the student to be failing due to "
        "excessive absences and/or failure to make up work due to absences."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Any student who ceases to attend class without officially withdrawing through the Admissions Office is '
        'subject to a grade of "F." The student will receive a grade of "W" for the course if withdrawn before the '
        '"last day to drop with a "W," and an "F" if withdrawn after "the last day to drop with a "W." Three tardies '
        'may constitute an absence.'
    )
    
    # Absence Policy / Make-up Work
    add_heading_with_style(doc, 'Absence Policy / Make-up Work / Extra Credit Policy', level=2)
    doc.add_paragraph(
        "Cisco College recognizes that absences from class may occur due to illness, death or illness in the immediate "
        "family, observance of a religious holiday, or participation in a College-sponsored activity. (Absences due to "
        "participation in a College-sponsored activity must be authorized by the appropriate administrator.) When absences "
        "occur due to the above, the student is allowed to make up work missed; the instructor may require the work to be "
        "made up within two weeks. There will be NO extra credit offered for this class!"
    )
    
    # Calculator Policy
    add_heading_with_style(doc, 'Calculator Policy', level=2)
    doc.add_paragraph(
        "Scientific and graphing calculators are generally permitted for this course. Calculators with CAS (Computer "
        "Algebra System) technology are NOT permitted. Examples of these calculators include the HP 50g, the HP Prime, "
        "the TI-Nspire CAS, the TI-89, TI-92, and the Casio ClassPad. Cell phone calculators and computers are never "
        "allowed for calculators in this class."
    )
    
    # Student Conduct
    add_heading_with_style(doc, 'Student Conduct, Notices, and College Policies', level=2)
    doc.add_paragraph(
        "Student Conduct, Notices, and College Policies Students are expected to follow all classroom policies listed in "
        "the course syllabus. College-wide policies can be found in the Official Catalog and the Student Handbook. "
        "Inappropriate behavior in the classroom shall result, at a minimum, in a request to leave class. The Student "
        "Handbook contains a list of specific prohibitions."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Dual credit courses are more challenging than high school courses and expect students to complete work on par "
        "with any other college student, demonstrating maturity and openness to new and varied ideas. Student information, "
        "attendance, and performance/grades will only be discussed with the student."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Dual credit classes will be cancelled in the event of Cisco College or ISD campus closure. Make-up work may be "
        "required. ISD campus events cannot disrupt dual credit classes."
    )
    
    # Academic Integrity
    add_heading_with_style(doc, 'Academic Integrity', level=2)
    doc.add_paragraph(
        "It is the intent of Cisco College to foster a spirit of complete honesty and a high standard of integrity. The "
        "attempt of students to present as their own any work they have not honestly performed is regarded by the faculty "
        "and administration as a serious offense and subjects the offender to disciplinary action. Violations are reported; "
        "multiple violations are tracked. The Student Handbook contains a list of academic integrity definitions, violations, "
        "and disciplinary actions. (Faculty may add to notice.)"
    )
    
    # Cross-Listed Course Sections
    add_heading_with_style(doc, 'Cross-Listed Course Sections', level=2)
    doc.add_paragraph(
        "For reasons of pedagogy and course management, this course may be cross-listed with one or more other course "
        "sections on Canvas. Cross-listed course sections may interact."
    )
    
    doc.add_page_break()
    
    # ========== PAGE 8-9 ==========
    # Course Content
    add_heading_with_style(doc, 'Course Content', level=2)
    doc.add_paragraph(
        "College-level courses may include controversial, sensitive, and/or adult material. Students are expected to have "
        "the readiness for college-level rigor and content."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Dual Credit students must follow Cisco College procedures to receive accommodations. High school IEPs do not apply "
        "to college courses. The high school counselor can assist you in contacting the Disability Services Coordinator."
    )
    
    # Disability Services
    add_heading_with_style(doc, 'Disability Services/ADA Accommodations', level=2)
    doc.add_paragraph(
        "Cisco College provides appropriate accommodations to qualified students in accordance with the Rehabilitation Act "
        "of 1973 and the Americans with Disabilities (ADA) Act of 1990. Accommodations are made on a case-by-case basis. "
        "Students with special needs are encouraged to contact the Disability Services Coordinator as early as possible. "
        "Early notice is required to prepare for and provide special accommodations by the first week of class. It is the "
        "student's responsibility to provide the necessary documentation to the Disability Services Coordinator prior to "
        "receiving accommodations."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Dual Credit students should follow Cisco College procedures to receive accommodations as high school accommodations "
        "do not apply to college courses. The high school counselor can assist you in contacting the Disability Services Coordinator."
    )
    
    # Title IX
    add_heading_with_style(doc, 'Title IX', level=2)
    doc.add_paragraph(
        "The college prohibits Sexual Misconduct and is committed to the timely and fair resolution of Sexual Misconduct cases. "
        "The College encourages prompt reporting of all types of Sexual Misconduct. The College has defined Sexual Misconduct as "
        "any unwelcome conduct of a sexual nature. The following persons may be contacted regarding Title IX issues: Title IX "
        "Coordinator (254-442-5022), Dean of Students (254-442-5173), Provost (325-794-4401) or any counselor."
    )
    
    # Parenting Students
    add_heading_with_style(doc, 'Parenting Students', level=2)
    doc.add_paragraph(
        "As a parent, the Title IX office can assist you with reasonable accommodations necessary for your academic success. "
        "These accommodations may be provided while a student is pregnant, during any pre- or post-delivery complications, and "
        "while parenting or caretaking. If you are a parent or guardian of a child younger than 18 years of age or expecting a "
        "child, please complete the Pregnancy and Parenting Support Form. This form is your opportunity to notify Cisco College "
        "that you are a parenting or pregnant student and/or may need accommodation due to parenting or pregnancy-related issues. "
        "Please note that pregnancy and parenting statuses apply to both partners, regardless of sex or gender identity."
    )
    
    # HB 1508
    add_heading_with_style(doc, 'HB 1508', level=2)
    doc.add_paragraph(
        "For students in this course who may have a criminal background, please be advised that the background could keep you "
        "from being licensed by the State of Texas. If you have a question about your background and licensure, please speak "
        "with your faculty member or the department chair. You also have the right to request a criminal history evaluation "
        "letter from the applicable licensing agency."
    )
    
    # Student Technology Use
    add_heading_with_style(doc, 'Student Technology Use in Classroom', level=2)
    doc.add_paragraph(
        "The use of communication devices is prohibited during class. Exceptions to this policy may occur due to college-wide "
        "emergency notification or at the discretion of the instructor. In order to protect the privacy of other students and "
        "to encourage open discussion, covert digital recording is prohibited in the classroom and material from online classes "
        "may not be recorded, shared, or reposted publicly. Students are expected to follow the Student IT Policy as stated in "
        "the Student Handbook."
    )
    
    # Student Help and Resources
    add_heading_with_style(doc, 'Student Help and Resources', level=2)
    doc.add_paragraph(
        "Students are encouraged to utilize the Canvas and online learning resources provided on the Distance Education webpage "
        "and the Student Resources provided on Canvas. For Canvas assistance: online@cisco.edu or 325-794-4480. For assistance "
        "with college computers, software, and email: helpdesk@students.cisco.edu."
    )
    
    # Changes to Syllabus
    add_heading_with_style(doc, 'Changes to the Syllabus', level=2)
    doc.add_paragraph(
        "The schedule and procedures in this syllabus are subject to change if deemed appropriate by the instructor."
    )
    
    # Save the document
    output_path = r"G:\My Drive\Colab Notebooks\COSC1315\Syllabus\COSC_1315_Syllabus_RECREATED.docx"
    doc.save(output_path)
    print(f"✅ Syllabus successfully recreated at:\n{output_path}")
    print(f"\n📄 Document contains {len(doc.paragraphs)} paragraphs")
    print(f"📊 Document contains {len(doc.tables)} tables")

if __name__ == "__main__":
    main()
